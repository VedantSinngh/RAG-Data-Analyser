import os
import uuid
import shutil
import logging
import hashlib
import httpx
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Request
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from pydantic import BaseModel
import chromadb

from pypdf import PdfReader
from docx import Document as DocxDocument
import pandas as pd
import numpy as np
import json

from app.config import settings
from app.database import get_db
from app.models.user import User
from app.models.document import Document
from app.models.embedding import Embedding
from app.routers.auth import get_current_user

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/documents", tags=["Documents & RAG"])

# Initialize Chroma Persistent Client
chroma_client = chromadb.PersistentClient(path=settings.CHROMA_PERSIST_DIR)

def get_chroma_collection():
    return chroma_client.get_or_create_collection(
        name=settings.CHROMA_COLLECTION_NAME
    )

class QueryRequest(BaseModel):
    query: str
    limit: int = 5

async def get_embeddings(texts: list[str], hf_api_key: str | None = None) -> list[list[float]]:
    """
    Generate text embeddings for RAG indexing.
    
    Groq does not offer an embeddings API, so we use this fallback chain:
      1. HuggingFace Inference API (free tier) — sentence-transformers/all-MiniLM-L6-v2
      2. Local deterministic hash-based vectors (offline, always works)
         Consistent dim=384 matching all-MiniLM-L6-v2 for schema compatibility.
    """
    api_key = hf_api_key or settings.HUGGINGFACE_API_KEY
    # 1. Try HuggingFace free embedding API if configured
    if api_key and api_key.strip():
        try:
            headers = {
                "Authorization": f"Bearer {api_key.strip()}",
                "Content-Type": "application/json"
            }
            async with httpx.AsyncClient() as client:
                response = await client.post(
                    "https://api-inference.huggingface.co/pipeline/feature-extraction/sentence-transformers/all-MiniLM-L6-v2",
                    headers=headers,
                    json={"inputs": texts, "options": {"wait_for_model": True}},
                    timeout=60.0
                )
                if response.status_code == 200:
                    data = response.json()
                    # HF returns list of embeddings when input is a list
                    if isinstance(data, list) and len(data) == len(texts):
                        result = []
                        for item in data:
                            embedding = item if isinstance(item[0], float) else item[0]
                            result.append(embedding)
                        logger.info(f"Generated {len(result)} embeddings via HuggingFace API")
                        return result
                else:
                    logger.warning(f"HuggingFace embedding API failed: {response.text}")
        except Exception as e:
            logger.warning(f"HuggingFace embedding error, falling back to local hash: {e}")

    # 2. Local fallback: Deterministic hash-based vector (dim=384, matches all-MiniLM-L6-v2)
    logger.info("Using local offline deterministic hash embeddings (no API key configured)")
    dim = 384
    embeddings = []
    for text in texts:
        h = hashlib.sha256(text.encode('utf-8')).digest()
        vector = []
        for i in range(dim):
            byte_idx = (i * 3) % len(h)
            val = (h[byte_idx] / 127.5) - 1.0
            vector.append(val)
        embeddings.append(vector)
    return embeddings


async def call_groq_summary(filename: str, sample_text: str) -> dict:
    """Helper to query Groq for a structured summary of unstructured text documents."""
    api_key = settings.GROQ_API_KEY
    if not api_key or api_key == "your-groq-api-key-here":
        return {
            "summary": "Document indexed successfully in local database.",
            "key_topics": ["General Reading", "Indexed Passage"],
            "suggested_questions": [
                "What is the main topic of this document?",
                "Provide a summary of the key findings in the text.",
                "Detail any specific statistics or numbers mentioned.",
                "Explain the primary arguments or sections.",
                "How does this connect to other uploaded papers?"
            ]
        }
        
    try:
        prompt = (
            "Analyze the following document preview and return a JSON object with the following keys:\n"
            "1. 'summary': A 2-3 sentence overview of what the document is about.\n"
            "2. 'key_topics': A list of up to 4 main subjects, sections, or entities discussed.\n"
            "3. 'suggested_questions': A list of exactly 5 precise questions that a reader could ask about this document to learn more from it.\n\n"
            f"Document Filename: {filename}\n"
            f"Text Sample:\n{sample_text[:3000]}\n\n"
            "Return ONLY a valid JSON block, no markdown format other than raw JSON."
        )
        headers = {
            "Authorization": f"Bearer {api_key.strip()}",
            "Content-Type": "application/json"
        }
        async with httpx.AsyncClient() as client:
            response = await client.post(
                f"{settings.GROQ_BASE_URL}/chat/completions",
                headers=headers,
                json={
                    "model": settings.GROQ_MODEL,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.2,
                    "max_tokens": 1000
                },
                timeout=15.0
            )
            if response.status_code == 200:
                content = response.json()["choices"][0]["message"]["content"].strip()
                if content.startswith("```"):
                    content = content.replace("```json", "").replace("```", "").strip()
                return json.loads(content)
    except Exception as e:
        logger.warning(f"Error querying Groq summary: {e}")
        
    return {
        "summary": "Context indexed in search database. Groq inference fallback activated.",
        "key_topics": ["General Information", "Ingested Context"],
        "suggested_questions": [
            "Summarize the core message of this document.",
            "What are the main key takeaways?",
            "What specific sections are in this document?",
            "List any statistics or data facts mentioned in the file.",
            "Can you explain the main arguments in this paper?"
        ]
    }


async def generate_document_metadata_async(filename: str, file_path: str, extension: str, text_content: str = "") -> dict:
    """Auto-generate statistics, column structures, row counts, preview, missing values, anomalies, and LLM suggested questions."""
    meta = {
        "is_tabular": False,
        "suggested_questions": []
    }
    
    if extension in ["csv", "xlsx", "xls"]:
        try:
            if extension == "csv":
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            meta["is_tabular"] = True
            meta["row_count"] = int(df.shape[0])
            meta["col_count"] = int(df.shape[1])
            meta["columns"] = df.columns.tolist()
            
            # Map column types
            col_types = {}
            for col in df.columns:
                col_type = str(df[col].dtype)
                if "int" in col_type or "float" in col_type:
                    col_types[col] = "numeric"
                elif "datetime" in col_type or df[col].astype(str).str.match(r'^\d{4}-\d{2}-\d{2}').all():
                    col_types[col] = "datetime"
                elif "bool" in col_type:
                    col_types[col] = "boolean"
                else:
                    col_types[col] = "categorical"
            meta["column_types"] = col_types
            
            # Row Preview (first 5 rows)
            preview_rows = []
            for _, row in df.head(5).iterrows():
                row_dict = {}
                for col in df.columns:
                    val = row[col]
                    if pd.isnull(val):
                        row_dict[col] = ""
                    elif isinstance(val, (np.integer, int)):
                        row_dict[col] = int(val)
                    elif isinstance(val, (np.floating, float)):
                        row_dict[col] = float(val)
                    elif isinstance(val, (np.bool_, bool)):
                        row_dict[col] = bool(val)
                    else:
                        row_dict[col] = str(val)
                preview_rows.append(row_dict)
            meta["preview"] = preview_rows
            
            # Null value count
            meta["null_counts"] = {col: int(df[col].isnull().sum()) for col in df.columns}
            
            # Duplicate rows count
            meta["duplicate_count"] = int(df.duplicated().sum())
            
            # Unique counts for categorical columns
            categorical_cols = [col for col, t in col_types.items() if t == "categorical"]
            meta["unique_counts"] = {col: int(df[col].nunique()) for col in categorical_cols}
            
            # Summary statistics for numeric columns
            numeric_cols = [col for col, t in col_types.items() if t == "numeric"]
            summary = {}
            for col in numeric_cols:
                series = df[col].dropna()
                if not series.empty:
                    summary[col] = {
                        "mean": float(series.mean()),
                        "median": float(series.median()),
                        "min": float(series.min()),
                        "max": float(series.max()),
                        "std": float(series.std()) if len(series) > 1 else 0.0
                    }
            meta["summary_stats"] = summary
            
            # Anomalies
            anomalies = []
            for col in df.columns:
                null_pct = df[col].isnull().sum() / len(df) if len(df) > 0 else 0
                if null_pct > 0.3:
                    anomalies.append(f"Column '{col}' has a high percentage of missing values ({null_pct*100:.1f}%).")
            
            for col in numeric_cols:
                series = df[col].dropna()
                if len(series) >= 4:
                    q1 = series.quantile(0.25)
                    q3 = series.quantile(0.75)
                    iqr = q3 - q1
                    lower = q1 - 1.5 * iqr
                    upper = q3 + 1.5 * iqr
                    outliers = series[(series < lower) | (series > upper)]
                    if not outliers.empty:
                        anomalies.append(f"Column '{col}' contains {len(outliers)} potential outliers outside IQR range [{lower:.2f}, {upper:.2f}].")
            meta["anomalies"] = anomalies
            
            # Call Groq to generate suggested questions if key exists
            api_key = settings.GROQ_API_KEY
            if api_key and api_key != "your-groq-api-key-here":
                try:
                    summary_info = {
                        "filename": filename,
                        "columns": meta["columns"],
                        "column_types": meta["column_types"],
                        "row_count": meta["row_count"],
                        "null_counts": meta["null_counts"],
                        "anomalies": meta["anomalies"]
                    }
                    prompt = (
                        "Review the following database structure summary and return a JSON object with exactly one key:\n"
                        "'suggested_questions': A list of exactly 5 precise questions that a user can ask to analyze trends, correlations, or anomalies in this dataset using natural language.\n\n"
                        f"Database Summary:\n{summary_info}\n\n"
                        "Return ONLY a valid JSON block, no markdown format other than raw JSON."
                    )
                    headers = {
                        "Authorization": f"Bearer {api_key.strip()}",
                        "Content-Type": "application/json"
                    }
                    async with httpx.AsyncClient() as client:
                        response = await client.post(
                            f"{settings.GROQ_BASE_URL}/chat/completions",
                            headers=headers,
                            json={
                                "model": settings.GROQ_MODEL,
                                "messages": [{"role": "user", "content": prompt}],
                                "temperature": 0.2,
                                "max_tokens": 500
                            },
                            timeout=10.0
                        )
                        if response.status_code == 200:
                            content = response.json()["choices"][0]["message"]["content"].strip()
                            if content.startswith("```"):
                                content = content.replace("```json", "").replace("```", "").strip()
                            res_json = json.loads(content)
                            meta["suggested_questions"] = res_json.get("suggested_questions", [])[:5]
                except Exception as ex:
                    logger.warning(f"Error querying Groq for tabular suggested questions: {ex}")
            
            # Default questions fallback
            if not meta["suggested_questions"]:
                questions = []
                if len(numeric_cols) > 0:
                    questions.append(f"Show summary statistics and distribution for {numeric_cols[0]}.")
                if len(categorical_cols) > 0 and len(numeric_cols) > 0:
                    questions.append(f"Compare average {numeric_cols[0]} grouped by {categorical_cols[0]}.")
                if len(numeric_cols) > 1:
                    questions.append(f"Analyze correlation between {numeric_cols[0]} and {numeric_cols[1]}.")
                questions.append("Identify and clean any duplicate rows or outliers in the dataset.")
                if any("date" in col.lower() or "time" in col.lower() for col in df.columns) and len(numeric_cols) > 0:
                    date_col = next(col for col in df.columns if "date" in col.lower() or "time" in col.lower())
                    questions.append(f"Forecast the next 12 periods of {numeric_cols[0]} using the {date_col} timeline.")
                else:
                    questions.append("Explain the key insights and trends in this data.")
                meta["suggested_questions"] = questions[:5]
                
        except Exception as e:
            logger.error(f"Error generating tabular metadata: {e}")
            meta["error"] = f"Failed to generate tabular metadata profile: {str(e)}"
    else:
        # Unstructured text document (PDF, DOCX, TXT)
        meta["is_tabular"] = False
        meta["word_count"] = len(text_content.split())
        meta["char_count"] = len(text_content)
        
        # Call Groq to get summary, key topics, suggested questions
        groq_res = await call_groq_summary(filename, text_content)
        meta["summary"] = groq_res.get("summary", "")
        meta["key_topics"] = groq_res.get("key_topics", [])
        meta["suggested_questions"] = groq_res.get("suggested_questions", [])[:5]
        
    return meta


# Parsers
def parse_text_file(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def parse_pdf_file(file_path: str) -> str:
    reader = PdfReader(file_path)
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return "\n\n".join(text_parts)

def parse_docx_file(file_path: str) -> str:
    doc = DocxDocument(file_path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    return "\n".join(text_parts)

def parse_csv_file(file_path: str) -> str:
    df = pd.read_csv(file_path)
    lines = []
    for idx, row in df.iterrows():
        row_str = ", ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
        lines.append(f"Row {idx + 1}: {row_str}")
    return "\n".join(lines)

def parse_excel_file(file_path: str) -> str:
    df = pd.read_excel(file_path)
    lines = []
    for idx, row in df.iterrows():
        row_str = ", ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
        lines.append(f"Row {idx + 1}: {row_str}")
    return "\n".join(lines)

def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> list[str]:
    chunks = []
    words = text.split()
    if not words:
        return []
    
    # 1000 chars roughly maps to 160 words
    words_per_chunk = max(10, chunk_size // 6)
    overlap_words = max(2, chunk_overlap // 6)
    
    step = words_per_chunk - overlap_words
    if step <= 0:
        step = words_per_chunk
        
    for i in range(0, len(words), step):
        chunk = " ".join(words[i:i + words_per_chunk])
        if chunk.strip():
            chunks.append(chunk)
    return chunks

@router.post("/upload", status_code=status.HTTP_201_CREATED)
async def upload_document(
    request: Request,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    uploads_dir = "/app/uploads"
    os.makedirs(uploads_dir, exist_ok=True)
    
    file_id = uuid.uuid4()
    extension = file.filename.split(".")[-1].lower() if "." in file.filename else ""
    local_filename = f"{file_id}.{extension}"
    file_path = os.path.join(uploads_dir, local_filename)
    
    try:
        with open(file_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
    except Exception as e:
        logger.error(f"Error saving uploaded file: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to write file to disk: {str(e)}"
        )
        
    size_bytes = os.path.getsize(file_path)
    
    text_content = ""
    try:
        if extension == "txt":
            text_content = parse_text_file(file_path)
        elif extension == "pdf":
            text_content = parse_pdf_file(file_path)
        elif extension in ["docx", "doc"]:
            text_content = parse_docx_file(file_path)
        elif extension == "csv":
            text_content = parse_csv_file(file_path)
        elif extension in ["xlsx", "xls"]:
            text_content = parse_excel_file(file_path)
        else:
            text_content = parse_text_file(file_path)
    except Exception as e:
        logger.error(f"Error parsing file content: {e}")
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to parse file formats: {str(e)}"
        )
        
    if not text_content.strip():
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Parsed file content is empty."
        )
        
    chunks = chunk_text(text_content)
    if not chunks:
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to segment text content into chunks."
        )
        
    hf_api_key = request.headers.get("X-HuggingFace-Api-Key")
    try:
        embeddings = await get_embeddings(chunks, hf_api_key=hf_api_key)
    except Exception as e:
        logger.error(f"Error calculating embeddings: {e}")
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to calculate embeddings: {str(e)}"
        )
        
    try:
        metadata_profile = await generate_document_metadata_async(
            file.filename, file_path, extension, text_content
        )
    except Exception as meta_err:
        logger.warning(f"Failed to generate metadata profile: {meta_err}")
        metadata_profile = {}
        
    metadata_profile["chunks_count"] = len(chunks)

    new_doc = Document(
        id=file_id,
        user_id=current_user.id,
        filename=file.filename,
        file_type=extension,
        storage_url=file_path,
        size_bytes=size_bytes,
        status="ready",
        metadata_json=metadata_profile
    )
    db.add(new_doc)
    
    chroma_ids = []
    chroma_embeddings = []
    chroma_metadatas = []
    chroma_texts = []
    
    for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
        c_id = str(uuid.uuid4())
        
        db_emb = Embedding(
            document_id=file_id,
            chunk_index=idx,
            chunk_text=chunk,
            chroma_id=c_id
        )
        db.add(db_emb)
        
        chroma_ids.append(c_id)
        chroma_embeddings.append(embedding)
        chroma_metadatas.append({
            "document_id": str(file_id),
            "user_id": str(current_user.id),
            "filename": file.filename,
            "chunk_index": idx
        })
        chroma_texts.append(chunk)
        
    try:
        get_chroma_collection().add(
            ids=chroma_ids,
            embeddings=chroma_embeddings,
            metadatas=chroma_metadatas,
            documents=chroma_texts
        )
    except Exception as e:
        logger.error(f"Error adding to ChromaDB: {e}")
        if "expecting embedding with dimension of" in str(e):
            logger.warning("Dimension mismatch detected in ChromaDB collection. Re-creating collection...")
            try:
                chroma_client.delete_collection(name=settings.CHROMA_COLLECTION_NAME)
                get_chroma_collection().add(
                    ids=chroma_ids,
                    embeddings=chroma_embeddings,
                    metadatas=chroma_metadatas,
                    documents=chroma_texts
                )
            except Exception as retry_err:
                logger.error(f"Failed to retry ChromaDB insertion after recreating collection: {retry_err}")
                if os.path.exists(file_path):
                    os.remove(file_path)
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"ChromaDB insert failed after dimension mismatch reset: {str(retry_err)}"
                )
        else:
            if os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"ChromaDB insert failed: {str(e)}"
            )
        
    await db.commit()
    await db.refresh(new_doc)
    
    return {
        "message": "Document uploaded and indexed successfully.",
        "id": str(new_doc.id),
        "filename": new_doc.filename,
        "size_bytes": new_doc.size_bytes,
        "chunks": len(chunks)
    }

@router.get("")
async def list_documents(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    result = await db.execute(
        select(Document)
        .filter(Document.user_id == current_user.id)
        .order_by(Document.created_at.desc())
    )
    docs = result.scalars().all()
    return [
        {
            "id": str(doc.id),
            "filename": doc.filename,
            "file_type": doc.file_type,
            "size_bytes": doc.size_bytes,
            "status": doc.status,
            "metadata": doc.metadata_json,
            "created_at": doc.created_at
        }
        for doc in docs
    ]

@router.delete("/{document_id}")
async def delete_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    try:
        doc_uuid = uuid.UUID(document_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid document ID format.")
        
    result = await db.execute(
        select(Document).filter(Document.id == doc_uuid, Document.user_id == current_user.id)
    )
    doc = result.scalars().first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found.")
        
    emb_result = await db.execute(
        select(Embedding).filter(Embedding.document_id == doc_uuid)
    )
    embeddings = emb_result.scalars().all()
    chroma_ids = [e.chroma_id for e in embeddings]
    
    if chroma_ids:
        try:
            get_chroma_collection().delete(ids=chroma_ids)
        except Exception as e:
            logger.error(f"Error deleting from ChromaDB: {e}")
            
    if doc.storage_url and os.path.exists(doc.storage_url):
        try:
            os.remove(doc.storage_url)
        except Exception as e:
            logger.error(f"Error removing local file from storage: {e}")
            
    await db.delete(doc)
    await db.commit()
    
    return {"message": "Document successfully deleted."}

@router.post("/query")
async def query_documents(
    req: QueryRequest,
    request: Request,
    current_user: User = Depends(get_current_user)
):
    if not req.query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty.")
        
    hf_api_key = request.headers.get("X-HuggingFace-Api-Key")
    try:
        query_embeddings = await get_embeddings([req.query], hf_api_key=hf_api_key)
        query_embedding = query_embeddings[0]
    except Exception as e:
        logger.error(f"Error embedding query: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate query vector: {str(e)}"
        )
        
    try:
        results = get_chroma_collection().query(
            query_embeddings=[query_embedding],
            n_results=req.limit,
            where={"user_id": str(current_user.id)}
        )
    except Exception as e:
        logger.error(f"Error querying ChromaDB: {e}")
        if "expecting embedding with dimension of" in str(e):
            logger.warning("Dimension mismatch in ChromaDB query. Re-creating collection...")
            try:
                chroma_client.delete_collection(name=settings.CHROMA_COLLECTION_NAME)
                get_chroma_collection()
                return {"results": []}
            except Exception:
                pass
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"ChromaDB search query failed: {str(e)}"
        )
        
    items = []
    if results and "documents" in results and results["documents"]:
        docs_list = results["documents"][0]
        ids_list = results["ids"][0]
        metadatas_list = results["metadatas"][0]
        
        # Resolve distances
        distances = (
            results["distances"][0] 
            if "distances" in results and results["distances"] 
            else [0.0] * len(docs_list)
        )
        
        for text, c_id, meta, dist in zip(docs_list, ids_list, metadatas_list, distances):
            confidence = max(0.0, min(1.0, 1.0 - (dist / 2.0)))
            items.append({
                "chroma_id": c_id,
                "text": text,
                "document_id": meta.get("document_id"),
                "filename": meta.get("filename"),
                "chunk_index": meta.get("chunk_index"),
                "confidence": round(confidence, 4)
            })
            
    return {"results": items}
